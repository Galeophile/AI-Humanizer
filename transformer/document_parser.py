"""
Document parser for .docx files with comprehensive formatting preservation.

This module provides data classes and functions for parsing Microsoft Word documents
while preserving formatting information including fonts, colors, styles, and structure.
"""

from dataclasses import dataclass
from typing import List, Optional, Union, Dict, Any
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_COLOR_TYPE
from docx.enum.dml import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import html
import pdfplumber
import re
from collections import defaultdict


@dataclass
class FormattedRun:
    """
    Represents a run of text with consistent formatting.
    
    Attributes:
        text: The text content of the run
        bold: Whether the text is bold
        italic: Whether the text is italic
        underline: Whether the text is underlined
        underline_style: The specific underline style if any
        font_name: The font family name
        font_size: The font size in points
        color: The text color in hex format (e.g., '#FF0000')
        highlight: The highlight color in hex format
    """
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    underline_style: Optional[str] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    color: Optional[str] = None
    highlight: Optional[str] = None


@dataclass
class FormattedParagraph:
    """
    Represents a paragraph with its formatting and runs.
    
    Attributes:
        runs: List of formatted runs in the paragraph
        style: The paragraph style name
        alignment: Text alignment as string
        space_before: Space before paragraph in points
        space_after: Space after paragraph in points
        line_spacing: Line spacing value
        left_indent: Left indent in points
        right_indent: Right indent in points
        first_line_indent: First line indent in points
    """
    runs: List[FormattedRun]
    style: Optional[str] = None
    alignment: Optional[str] = None
    space_before: Optional[float] = None
    space_after: Optional[float] = None
    line_spacing: Optional[float] = None
    left_indent: Optional[float] = None
    right_indent: Optional[float] = None
    first_line_indent: Optional[float] = None


@dataclass
class FormattedListItem:
    """
    Represents a list item with its level and type.
    
    Attributes:
        paragraph: The formatted paragraph content
        level: The nesting level (0-based)
        list_type: The type of list ('bullet', 'number', etc.)
        number_format: The numbering format if applicable
    """
    paragraph: FormattedParagraph
    level: int
    list_type: str
    number_format: Optional[str] = None


@dataclass
class FormattedTable:
    """
    Represents a table with its cell contents.
    
    Attributes:
        rows: List of rows, each containing a list of cell paragraphs
        style: The table style name
    """
    rows: List[List[List[FormattedParagraph]]]
    style: Optional[str] = None


@dataclass
class FormattedDocument:
    """
    Represents a complete document with all its content.
    
    Attributes:
        paragraphs: List of formatted paragraphs
        tables: List of formatted tables
        list_items: List of formatted list items
        styles: Document styles information
    """
    paragraphs: List[FormattedParagraph]
    tables: List[FormattedTable]
    list_items: List[FormattedListItem]
    styles: Optional[Dict[str, Any]] = None


def rgb_to_hex(rgb) -> Optional[str]:
    """
    Convert RGB color to hex format.
    
    Args:
        rgb: RGB color object or tuple
        
    Returns:
        Hex color string (e.g., '#FF0000') or None if conversion fails
    """
    try:
        if hasattr(rgb, 'rgb'):
            color = rgb.rgb
            if color is not None:
                return f"#{color:06X}"
        elif isinstance(rgb, tuple) and len(rgb) == 3:
            return f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
        elif isinstance(rgb, RGBColor):
            return f"#{rgb.rgb:06X}"
    except (AttributeError, TypeError, ValueError):
        pass
    return None


def alignment_to_string(alignment) -> Optional[str]:
    """
    Convert alignment enum to string representation.
    
    Args:
        alignment: WD_ALIGN_PARAGRAPH enum value
        
    Returns:
        String representation of alignment or None
    """
    try:
        if alignment is None:
            return None
        
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute"
        }
        return alignment_map.get(alignment)
    except (AttributeError, TypeError):
        return None


def extract_run_formatting(run) -> FormattedRun:
    """
    Extract formatting information from a document run.
    
    Args:
        run: The document run object
        
    Returns:
        FormattedRun object with extracted formatting
    """
    try:
        text = run.text or ""
        bold = run.bold or False
        italic = run.italic or False
        
        # Handle underline with style detection
        underline = False
        underline_style = None
        if hasattr(run, 'underline') and run.underline is not None:
            if run.underline is True or (hasattr(run.underline, 'val') and run.underline.val):
                underline = True
                if hasattr(run.underline, 'val') and run.underline.val != True:
                    underline_style = str(run.underline.val)
        
        # Extract font information
        font_name = None
        font_size = None
        if hasattr(run, 'font'):
            if hasattr(run.font, 'name') and run.font.name:
                font_name = run.font.name
            if hasattr(run.font, 'size') and run.font.size:
                # Convert from Emu to points (1 point = 12700 EMU)
                font_size = float(run.font.size.pt) if hasattr(run.font.size, 'pt') else None
        
        # Extract color information
        color = None
        if hasattr(run, 'font') and hasattr(run.font, 'color'):
            font_color = run.font.color
            if hasattr(font_color, 'rgb') and font_color.rgb is not None:
                color = rgb_to_hex(font_color.rgb)
            elif hasattr(font_color, 'theme_color') and font_color.theme_color is not None:
                # Handle theme colors - would need document theme for full resolution
                pass
        
        # Extract highlight information
        highlight = None
        if hasattr(run, 'font') and hasattr(run.font, 'highlight_color'):
            highlight_color = run.font.highlight_color
            if highlight_color is not None:
                # Store the enum name for round-trip fidelity
                if hasattr(highlight_color, 'name'):
                    highlight = highlight_color.name
                elif hasattr(highlight_color, 'rgb') and highlight_color.rgb is not None:
                    highlight = rgb_to_hex(highlight_color.rgb)
        
        return FormattedRun(
            text=text,
            bold=bold,
            italic=italic,
            underline=underline,
            underline_style=underline_style,
            font_name=font_name,
            font_size=font_size,
            color=color,
            highlight=highlight
        )
    except Exception:
        # Return basic run with just text if formatting extraction fails
        return FormattedRun(text=getattr(run, 'text', ''))


def extract_paragraph_formatting(paragraph) -> FormattedParagraph:
    """
    Extract formatting information from a document paragraph.
    
    Args:
        paragraph: The document paragraph object
        
    Returns:
        FormattedParagraph object with extracted formatting
    """
    try:
        # Extract runs
        runs = [extract_run_formatting(run) for run in paragraph.runs]
        
        # Extract paragraph style
        style = None
        if hasattr(paragraph, 'style') and paragraph.style:
            style = paragraph.style.name
        
        # Extract alignment
        alignment = alignment_to_string(paragraph.alignment)
        
        # Extract spacing and indent information
        space_before = None
        space_after = None
        line_spacing = None
        left_indent = None
        right_indent = None
        first_line_indent = None
        
        if hasattr(paragraph, 'paragraph_format'):
            pf = paragraph.paragraph_format
            
            if hasattr(pf, 'space_before') and pf.space_before:
                space_before = float(pf.space_before.pt) if hasattr(pf.space_before, 'pt') else None
            
            if hasattr(pf, 'space_after') and pf.space_after:
                space_after = float(pf.space_after.pt) if hasattr(pf.space_after, 'pt') else None
            
            if hasattr(pf, 'line_spacing') and pf.line_spacing:
                line_spacing = float(pf.line_spacing)
            
            if hasattr(pf, 'left_indent') and pf.left_indent:
                left_indent = float(pf.left_indent.pt) if hasattr(pf.left_indent, 'pt') else None
            
            if hasattr(pf, 'right_indent') and pf.right_indent:
                right_indent = float(pf.right_indent.pt) if hasattr(pf.right_indent, 'pt') else None
            
            if hasattr(pf, 'first_line_indent') and pf.first_line_indent:
                first_line_indent = float(pf.first_line_indent.pt) if hasattr(pf.first_line_indent, 'pt') else None
        
        return FormattedParagraph(
            runs=runs,
            style=style,
            alignment=alignment,
            space_before=space_before,
            space_after=space_after,
            line_spacing=line_spacing,
            left_indent=left_indent,
            right_indent=right_indent,
            first_line_indent=first_line_indent
        )
    except Exception:
        # Return basic paragraph with just runs if formatting extraction fails
        runs = []
        try:
            runs = [extract_run_formatting(run) for run in paragraph.runs]
        except:
            pass
        return FormattedParagraph(runs=runs)


def detect_list_item(paragraph) -> Optional[tuple]:
    """
    Detect if a paragraph is a list item and extract list information.
    
    Args:
        paragraph: The document paragraph object
        
    Returns:
        Tuple of (level, list_type, number_format) or None if not a list item
    """
    try:
        # Check for numbering properties using xpath
        num_pr = paragraph._element.xpath('.//w:numPr')
        if not num_pr:
            return None
        
        level = 0
        list_type = "bullet"
        number_format = None
        
        # Extract numbering level
        level_elements = paragraph._element.xpath('.//w:numPr/w:ilvl')
        if level_elements:
            try:
                level = int(level_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0'))
            except (ValueError, AttributeError):
                level = 0
        
        # Extract numbering ID to determine list type
        num_id_elements = paragraph._element.xpath('.//w:numPr/w:numId')
        if num_id_elements and hasattr(paragraph, '_parent') and hasattr(paragraph._parent, 'part'):
            try:
                num_id = num_id_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if num_id and hasattr(paragraph._parent.part, 'numbering_definitions'):
                    # Try to determine list type from numbering definitions
                    numbering_defs = paragraph._parent.part.numbering_definitions
                    if hasattr(numbering_defs, 'num_having_numId'):
                        num_def = numbering_defs.num_having_numId(int(num_id))
                        if num_def and hasattr(num_def, 'abstractNum'):
                            abstract_num = num_def.abstractNum
                            if hasattr(abstract_num, 'lvl_lst'):
                                lvl_list = abstract_num.lvl_lst
                                if len(lvl_list) > level:
                                    lvl_def = lvl_list[level]
                                    if hasattr(lvl_def, 'numFmt') and lvl_def.numFmt:
                                        fmt_val = lvl_def.numFmt.val
                                        if fmt_val == 'bullet':
                                            list_type = "bullet"
                                        elif fmt_val in ['decimal', 'upperRoman', 'lowerRoman', 'upperLetter', 'lowerLetter']:
                                            list_type = "number"
                                            number_format = fmt_val
                                        else:
                                            list_type = "number"
                                            number_format = fmt_val
            except (ValueError, AttributeError, IndexError):
                pass
        
        return (level, list_type, number_format)
    except Exception:
        return None


def extract_table_formatting(table) -> FormattedTable:
    """
    Extract formatting information from a document table.
    
    Args:
        table: The document table object
        
    Returns:
        FormattedTable object with extracted formatting
    """
    try:
        rows = []
        for table_row in table.rows:
            row_cells = []
            for cell in table_row.cells:
                cell_paragraphs = []
                for paragraph in cell.paragraphs:
                    cell_paragraphs.append(extract_paragraph_formatting(paragraph))
                row_cells.append(cell_paragraphs)
            rows.append(row_cells)
        
        # Extract table style
        style = None
        if hasattr(table, 'style') and table.style:
            style = table.style.name
        
        return FormattedTable(rows=rows, style=style)
    except Exception:
        return FormattedTable(rows=[])


def parse_docx_file(file_or_bytes: Union[str, bytes, io.BytesIO]) -> FormattedDocument:
    """
    Parse a .docx file and build a FormattedDocument.
    
    Args:
        file_or_bytes: File path, bytes, or BytesIO object containing the .docx file
        
    Returns:
        FormattedDocument object containing all document content with formatting
        
    Raises:
        Exception: If document parsing fails
    """
    try:
        # Load the document
        if isinstance(file_or_bytes, (str, bytes)):
            doc = Document(file_or_bytes)
        else:
            doc = Document(file_or_bytes)
        
        paragraphs = []
        tables = []
        list_items = []
        
        # Process document elements in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        formatted_para = extract_paragraph_formatting(para)
                        
                        # Check if it's a list item
                        list_info = detect_list_item(para)
                        if list_info:
                            level, list_type, number_format = list_info
                            list_item = FormattedListItem(
                                paragraph=formatted_para,
                                level=level,
                                list_type=list_type,
                                number_format=number_format
                            )
                            list_items.append(list_item)
                        else:
                            paragraphs.append(formatted_para)
                        break
            
            elif element.tag.endswith('tbl'):  # Table
                # Find corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        formatted_table = extract_table_formatting(table)
                        tables.append(formatted_table)
                        break
        
        # Extract document styles information
        styles = {}
        if hasattr(doc, 'styles'):
            try:
                for style in doc.styles:
                    if hasattr(style, 'name') and style.name:
                        styles[style.name] = {
                            'type': getattr(style, 'type', None),
                            'builtin': getattr(style, 'builtin', None)
                        }
            except Exception:
                pass
        
        return FormattedDocument(
            paragraphs=paragraphs,
            tables=tables,
            list_items=list_items,
            styles=styles if styles else None
        )
    
    except Exception as e:
        raise Exception(f"Failed to parse document: {str(e)}")


def get_plain_text(formatted_doc: FormattedDocument) -> str:
    """
    Extract plain text from a FormattedDocument.
    
    Args:
        formatted_doc: The FormattedDocument object
        
    Returns:
        Plain text string with paragraph breaks
    """
    text_parts = []
    
    # Add paragraph text
    for paragraph in formatted_doc.paragraphs:
        paragraph_text = ''.join(run.text for run in paragraph.runs)
        if paragraph_text.strip():
            text_parts.append(paragraph_text)
    
    # Add list item text
    for list_item in formatted_doc.list_items:
        paragraph_text = ''.join(run.text for run in list_item.paragraph.runs)
        if paragraph_text.strip():
            text_parts.append(paragraph_text)
    
    # Add table text
    for table in formatted_doc.tables:
        for row in table.rows:
            for cell in row:
                for paragraph in cell:
                    paragraph_text = ''.join(run.text for run in paragraph.runs)
                    if paragraph_text.strip():
                        text_parts.append(paragraph_text)
    
    return '\n'.join(text_parts)


def merge_consecutive_runs(runs: List[FormattedRun]) -> List[FormattedRun]:
    """
    Merge consecutive runs with identical formatting.
    
    Args:
        runs: List of FormattedRun objects
        
    Returns:
        List of merged FormattedRun objects
    """
    if not runs:
        return []
    
    merged = []
    current_run = runs[0]
    
    for next_run in runs[1:]:
        # Check if formatting is identical (excluding text)
        if (current_run.bold == next_run.bold and
            current_run.italic == next_run.italic and
            current_run.underline == next_run.underline and
            current_run.underline_style == next_run.underline_style and
            current_run.font_name == next_run.font_name and
            current_run.font_size == next_run.font_size and
            current_run.color == next_run.color and
            current_run.highlight == next_run.highlight):
            # Merge text
            current_run = FormattedRun(
                text=current_run.text + next_run.text,
                bold=current_run.bold,
                italic=current_run.italic,
                underline=current_run.underline,
                underline_style=current_run.underline_style,
                font_name=current_run.font_name,
                font_size=current_run.font_size,
                color=current_run.color,
                highlight=current_run.highlight
            )
        else:
            # Different formatting, add current and start new
            merged.append(current_run)
            current_run = next_run
    
    # Add the last run
    merged.append(current_run)
    return merged


def escape_html_text(text: str) -> str:
    """
    Escape HTML special characters in text.
    
    Args:
        text: Raw text to escape
        
    Returns:
        HTML-escaped text
    """
    return html.escape(text)


def formatted_run_to_html(run: FormattedRun) -> str:
    """
    Convert a FormattedRun to HTML with inline styles.
    
    Args:
        run: FormattedRun object to convert
        
    Returns:
        HTML string representing the run
    """
    if not run.text:
        return ""
    
    escaped_text = escape_html_text(run.text)
    styles = []
    
    # Font styling
    if run.bold:
        styles.append("font-weight: bold")
    if run.italic:
        styles.append("font-style: italic")
    if run.underline:
        if run.underline_style:
            styles.append("text-decoration: underline")
        else:
            styles.append("text-decoration: underline")
    if run.font_name:
        styles.append(f"font-family: '{run.font_name}'")
    if run.font_size:
        styles.append(f"font-size: {run.font_size}pt")
    if run.color:
        styles.append(f"color: {run.color}")
    if run.highlight:
        styles.append(f"background-color: {run.highlight}")
    
    if styles:
        style_attr = f' style="{"; ".join(styles)}"'
        return f"<span{style_attr}>{escaped_text}</span>"
    else:
        return escaped_text


def formatted_paragraph_to_html(paragraph: FormattedParagraph) -> str:
    """
    Convert a FormattedParagraph to HTML with inline styles.
    
    Args:
        paragraph: FormattedParagraph object to convert
        
    Returns:
        HTML string representing the paragraph
    """
    if not paragraph.runs:
        return "<p></p>"
    
    # Convert runs to HTML
    run_html = "".join(formatted_run_to_html(run) for run in paragraph.runs)
    
    # Build paragraph styles
    styles = []
    if paragraph.alignment:
        styles.append(f"text-align: {paragraph.alignment}")
    if paragraph.space_before:
        styles.append(f"margin-top: {paragraph.space_before}pt")
    if paragraph.space_after:
        styles.append(f"margin-bottom: {paragraph.space_after}pt")
    if paragraph.line_spacing:
        styles.append(f"line-height: {paragraph.line_spacing}")
    if paragraph.left_indent:
        styles.append(f"margin-left: {paragraph.left_indent}pt")
    if paragraph.right_indent:
        styles.append(f"margin-right: {paragraph.right_indent}pt")
    if paragraph.first_line_indent:
        styles.append(f"text-indent: {paragraph.first_line_indent}pt")
    
    style_attr = f' style="{"; ".join(styles)}"' if styles else ""
    return f"<p{style_attr}>{run_html}</p>"


def formatted_list_items_to_html(list_items: List[FormattedListItem]) -> str:
    """
    Convert list items to HTML with proper nesting.
    
    Args:
        list_items: List of FormattedListItem objects
        
    Returns:
        HTML string representing the nested lists
    """
    if not list_items:
        return ""
    
    html_parts = []
    stack = []  # Stack to track open list tags
    current_level = -1
    
    for item in list_items:
        # Close lists if level decreased
        while current_level >= item.level and stack:
            tag = stack.pop()
            html_parts.append(f"</{tag}>")
            current_level -= 1
        
        # Open new lists if level increased
        while current_level < item.level:
            current_level += 1
            if item.list_type == "bullet":
                tag = "ul"
            else:
                tag = "ol"
                if item.number_format:
                    if item.number_format == "upperRoman":
                        html_parts.append(f'<{tag} style="list-style-type: upper-roman">')
                    elif item.number_format == "lowerRoman":
                        html_parts.append(f'<{tag} style="list-style-type: lower-roman">')
                    elif item.number_format == "upperLetter":
                        html_parts.append(f'<{tag} style="list-style-type: upper-alpha">')
                    elif item.number_format == "lowerLetter":
                        html_parts.append(f'<{tag} style="list-style-type: lower-alpha">')
                    else:
                        html_parts.append(f"<{tag}>")
                else:
                    html_parts.append(f"<{tag}>")
            stack.append(tag)
        
        # Add list item
        paragraph_html = formatted_paragraph_to_html(item.paragraph)
        # Remove <p> tags for list items and use content directly
        content = paragraph_html.replace("<p>", "").replace("</p>", "").replace('<p style="', '<span style="').replace("</p>", "</span>")
        html_parts.append(f"<li>{content}</li>")
    
    # Close remaining open lists
    while stack:
        tag = stack.pop()
        html_parts.append(f"</{tag}>")
    
    return "".join(html_parts)


def formatted_table_to_html(table: FormattedTable) -> str:
    """
    Convert a FormattedTable to HTML.
    
    Args:
        table: FormattedTable object to convert
        
    Returns:
        HTML string representing the table
    """
    if not table.rows:
        return ""
    
    html_parts = ['<table style="border-collapse: collapse; width: 100%;">']
    
    for row in table.rows:
        html_parts.append("<tr>")
        for cell in row:
            html_parts.append('<td style="border: 1px solid #ddd; padding: 8px;">')
            for paragraph in cell:
                html_parts.append(formatted_paragraph_to_html(paragraph))
            html_parts.append("</td>")
        html_parts.append("</tr>")
    
    html_parts.append("</table>")
    return "".join(html_parts)


def formatted_document_to_html(document: FormattedDocument) -> str:
    """
    Convert a FormattedDocument to HTML.
    
    Args:
        document: FormattedDocument object to convert
        
    Returns:
        Complete HTML string representing the document
    """
    html_parts = []
    
    # Combine all elements in order they appear
    all_elements = []
    
    # Add paragraphs
    for paragraph in document.paragraphs:
        all_elements.append(("paragraph", paragraph))
    
    # Add list items
    if document.list_items:
        all_elements.append(("list", document.list_items))
    
    # Add tables
    for table in document.tables:
        all_elements.append(("table", table))
    
    # For now, render in the order they were stored
    # Regular paragraphs
    for paragraph in document.paragraphs:
        html_parts.append(formatted_paragraph_to_html(paragraph))
    
    # List items
    if document.list_items:
        html_parts.append(formatted_list_items_to_html(document.list_items))
    
    # Tables
    for table in document.tables:
        html_parts.append(formatted_table_to_html(table))
    
    return "".join(html_parts)


def add_formatted_run_to_paragraph(docx_paragraph, run: FormattedRun):
    """
    Add a formatted run to a python-docx paragraph.
    
    Args:
        docx_paragraph: python-docx paragraph object
        run: FormattedRun object to add
    """
    if not run.text:
        return
    
    docx_run = docx_paragraph.add_run(run.text)
    
    # Apply formatting
    if run.bold:
        docx_run.bold = True
    if run.italic:
        docx_run.italic = True
    if run.underline:
        docx_run.underline = True
    
    # Apply font formatting
    if run.font_name:
        docx_run.font.name = run.font_name
    if run.font_size:
        docx_run.font.size = Pt(run.font_size)
    
    # Apply color
    if run.color and run.color.startswith('#'):
        try:
            # Convert hex to RGB
            hex_color = run.color[1:]
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            docx_run.font.color.rgb = RGBColor(*rgb)
        except (ValueError, IndexError):
            pass
    
    # Apply highlight color
    if run.highlight:
        try:
            # Try to map back to WD_COLOR_INDEX enum
            highlight_map = {
                'YELLOW': WD_COLOR_INDEX.YELLOW,
                'BRIGHT_GREEN': WD_COLOR_INDEX.BRIGHT_GREEN,
                'TURQUOISE': WD_COLOR_INDEX.TURQUOISE,
                'PINK': WD_COLOR_INDEX.PINK,
                'BLUE': WD_COLOR_INDEX.BLUE,
                'RED': WD_COLOR_INDEX.RED,
                'DARK_BLUE': WD_COLOR_INDEX.DARK_BLUE,
                'TEAL': WD_COLOR_INDEX.TEAL,
                'GREEN': WD_COLOR_INDEX.GREEN,
                'VIOLET': WD_COLOR_INDEX.VIOLET,
                'DARK_RED': WD_COLOR_INDEX.DARK_RED,
                'DARK_YELLOW': WD_COLOR_INDEX.DARK_YELLOW,
                'GRAY_25': WD_COLOR_INDEX.GRAY_25,
                'GRAY_50': WD_COLOR_INDEX.GRAY_50,
                'WHITE': WD_COLOR_INDEX.WHITE,
                'BLACK': WD_COLOR_INDEX.BLACK
            }
            if run.highlight in highlight_map:
                docx_run.font.highlight_color = highlight_map[run.highlight]
        except (AttributeError, KeyError):
            pass


def add_formatted_paragraph_to_document(doc, paragraph: FormattedParagraph):
    """
    Add a formatted paragraph to a python-docx document.
    
    Args:
        doc: python-docx Document object
        paragraph: FormattedParagraph object to add
        
    Returns:
        The created paragraph object
    """
    docx_paragraph = doc.add_paragraph()
    
    # Add runs
    for run in paragraph.runs:
        add_formatted_run_to_paragraph(docx_paragraph, run)
    
    # Apply paragraph formatting
    if paragraph.alignment:
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE
        }
        if paragraph.alignment in alignment_map:
            docx_paragraph.alignment = alignment_map[paragraph.alignment]
    
    # Apply spacing and indent
    if paragraph.space_before:
        docx_paragraph.paragraph_format.space_before = Pt(paragraph.space_before)
    if paragraph.space_after:
        docx_paragraph.paragraph_format.space_after = Pt(paragraph.space_after)
    if paragraph.line_spacing:
        docx_paragraph.paragraph_format.line_spacing = paragraph.line_spacing
    if paragraph.left_indent:
        docx_paragraph.paragraph_format.left_indent = Pt(paragraph.left_indent)
    if paragraph.right_indent:
        docx_paragraph.paragraph_format.right_indent = Pt(paragraph.right_indent)
    if paragraph.first_line_indent:
        docx_paragraph.paragraph_format.first_line_indent = Pt(paragraph.first_line_indent)
    
    return docx_paragraph


def add_formatted_list_item_to_document(doc, list_item: FormattedListItem):
    """
    Add a formatted list item to a python-docx document.
    
    Args:
        doc: python-docx Document object
        list_item: FormattedListItem object to add
    """
    docx_paragraph = add_formatted_paragraph_to_document(doc, list_item.paragraph)
    
    # Create numbering properties using OxmlElement
    num_pr = OxmlElement('w:numPr')
    
    # Set numbering level
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(list_item.level))
    num_pr.append(ilvl)
    
    # Set numbering ID (use a simple mapping based on list type)
    num_id = OxmlElement('w:numId')
    if list_item.list_type == 'bullet':
        num_id.set(qn('w:val'), '1')  # Bullet list
    else:
        num_id.set(qn('w:val'), '2')  # Numbered list
    num_pr.append(num_id)
    
    # Apply numbering to paragraph
    docx_paragraph._element.get_or_add_pPr().append(num_pr)


def add_formatted_table_to_document(doc, table: FormattedTable):
    """
    Add a formatted table to a python-docx document.
    
    Args:
        doc: python-docx Document object
        table: FormattedTable object to add
        
    Returns:
        The created table object
    """
    if not table.rows:
        return None
    
    # Determine table dimensions
    max_cols = max(len(row) for row in table.rows) if table.rows else 0
    if max_cols == 0:
        return None
    
    # Create table
    docx_table = doc.add_table(rows=len(table.rows), cols=max_cols)
    
    # Apply table style if available
    if table.style:
        try:
            docx_table.style = table.style
        except (KeyError, ValueError):
            pass
    
    # Fill table cells
    for row_idx, row in enumerate(table.rows):
        docx_row = docx_table.rows[row_idx]
        for col_idx, cell_paragraphs in enumerate(row):
            if col_idx < len(docx_row.cells):
                docx_cell = docx_row.cells[col_idx]
                # Clear default paragraph
                docx_cell.paragraphs[0].clear()
                
                # Add paragraphs to cell
                for para_idx, paragraph in enumerate(cell_paragraphs):
                    if para_idx == 0:
                        # Use existing paragraph
                        cell_para = docx_cell.paragraphs[0]
                        for run in paragraph.runs:
                            add_formatted_run_to_paragraph(cell_para, run)
                    else:
                        # Add new paragraph
                        cell_para = docx_cell.add_paragraph()
                        for run in paragraph.runs:
                            add_formatted_run_to_paragraph(cell_para, run)
    
    return docx_table


def formatted_document_to_docx(document: FormattedDocument) -> bytes:
    """
    Convert a FormattedDocument to DOCX bytes.
    
    Args:
        document: FormattedDocument object to convert
        
    Returns:
        Bytes representing the DOCX document
    """
    doc = Document()
    
    # Add paragraphs
    for paragraph in document.paragraphs:
        add_formatted_paragraph_to_document(doc, paragraph)
    
    # Add list items
    for list_item in document.list_items:
        add_formatted_list_item_to_document(doc, list_item)
    
    # Add tables
    for table in document.tables:
        add_formatted_table_to_document(doc, table)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def infer_font_style_from_name(font_name: str) -> tuple[bool, bool]:
    """
    Infer bold and italic styles from font name.
    
    Args:
        font_name: The font name string
        
    Returns:
        Tuple of (is_bold, is_italic)
    """
    if not font_name:
        return False, False
    
    font_lower = font_name.lower()
    is_bold = any(keyword in font_lower for keyword in ['bold', 'black', 'heavy', 'extra'])
    is_italic = any(keyword in font_lower for keyword in ['italic', 'oblique', 'slant'])
    
    return is_bold, is_italic


def detect_bullet_or_number(text: str) -> tuple[str, int]:
    """
    Detect if text starts with bullet or number marker.
    
    Args:
        text: Text to analyze
        
    Returns:
        Tuple of (list_type, level) where level is inferred from indentation
    """
    if not text:
        return "none", 0
    
    stripped_text = text.lstrip()
    leading_spaces = len(text) - len(stripped_text)
    level = leading_spaces // 20  # Rough estimate, 20 chars per level
    
    # Check for bullet markers
    bullet_patterns = [r'^[•·▪▫■□▶▷◆◇○●★☆]\s+', r'^[-*+]\s+']
    for pattern in bullet_patterns:
        if re.match(pattern, stripped_text):
            return "bullet", level
    
    # Check for number markers
    number_patterns = [
        r'^\d+[.):]\s+',  # 1. 1) 1:
        r'^[a-zA-Z][.):]\s+',  # a. a) a:
        r'^[ivxlcdm]+[.):]\s+',  # roman numerals
        r'^[IVXLCDM]+[.):]\s+'   # ROMAN NUMERALS
    ]
    for pattern in number_patterns:
        if re.match(pattern, stripped_text):
            return "number", level
    
    return "none", 0


def group_text_into_paragraphs(chars: list) -> list:
    """
    Group character objects into paragraphs based on y-coordinates.
    
    Args:
        chars: List of character objects from pdfplumber
        
    Returns:
        List of character groups representing paragraphs
    """
    if not chars:
        return []
    
    # Sort characters by y-coordinate (top to bottom) then x-coordinate
    sorted_chars = sorted(chars, key=lambda c: (-c['y0'], c['x0']))
    
    paragraphs = []
    current_paragraph = []
    current_y = None
    line_height_threshold = 5  # Pixels
    
    for char in sorted_chars:
        if current_y is None:
            current_y = char['y0']
            current_paragraph = [char]
        elif abs(char['y0'] - current_y) > line_height_threshold:
            # New line/paragraph
            if current_paragraph:
                paragraphs.append(current_paragraph)
            current_paragraph = [char]
            current_y = char['y0']
        else:
            current_paragraph.append(char)
    
    if current_paragraph:
        paragraphs.append(current_paragraph)
    
    return paragraphs


def extract_text_from_char_group(char_group: list) -> str:
    """
    Extract text from a group of characters.
    
    Args:
        char_group: List of character objects
        
    Returns:
        Combined text string
    """
    if not char_group:
        return ""
    
    # Sort by x-coordinate to get proper reading order
    sorted_chars = sorted(char_group, key=lambda c: c['x0'])
    return ''.join(char.get('text', '') for char in sorted_chars)


def infer_formatting_from_chars(char_group: list) -> dict:
    """
    Infer formatting from character group properties.
    
    Args:
        char_group: List of character objects
        
    Returns:
        Dictionary with formatting properties
    """
    if not char_group:
        return {}
    
    # Use the first character as representative
    first_char = char_group[0]
    
    font_name = first_char.get('fontname', '')
    font_size = first_char.get('size', 12)
    
    is_bold, is_italic = infer_font_style_from_name(font_name)
    
    return {
        'font_name': font_name,
        'font_size': font_size,
        'bold': is_bold,
        'italic': is_italic,
        'color': None  # PDF color extraction is complex, keep simple for now
    }


def split_chars_by_formatting(char_group: list) -> list:
    """
    Split character group into runs with consistent formatting.
    
    Args:
        char_group: List of character objects
        
    Returns:
        List of character sub-groups with consistent formatting
    """
    if not char_group:
        return []
    
    runs = []
    current_run = []
    current_font = None
    current_size = None
    
    for char in char_group:
        char_font = char.get('fontname', '')
        char_size = char.get('size', 12)
        
        if current_font is None:
            current_font = char_font
            current_size = char_size
            current_run = [char]
        elif char_font == current_font and char_size == current_size:
            current_run.append(char)
        else:
            # Different formatting, start new run
            if current_run:
                runs.append(current_run)
            current_run = [char]
            current_font = char_font
            current_size = char_size
    
    if current_run:
        runs.append(current_run)
    
    return runs


def extract_paragraph_from_chars(char_group: list) -> FormattedParagraph:
    """
    Convert character group to FormattedParagraph.
    
    Args:
        char_group: List of character objects from pdfplumber
        
    Returns:
        FormattedParagraph object
    """
    if not char_group:
        return FormattedParagraph(runs=[])
    
    # Split into runs by formatting
    formatting_runs = split_chars_by_formatting(char_group)
    
    formatted_runs = []
    for run_chars in formatting_runs:
        text = extract_text_from_char_group(run_chars)
        if text.strip():  # Only create runs with content
            formatting = infer_formatting_from_chars(run_chars)
            
            formatted_run = FormattedRun(
                text=text,
                bold=formatting.get('bold', False),
                italic=formatting.get('italic', False),
                font_name=formatting.get('font_name'),
                font_size=formatting.get('font_size'),
                color=formatting.get('color')
            )
            formatted_runs.append(formatted_run)
    
    # Calculate basic paragraph properties from position
    left_indent = min(char['x0'] for char in char_group) if char_group else 0
    
    return FormattedParagraph(
        runs=formatted_runs,
        left_indent=left_indent if left_indent > 50 else None  # Only set if significantly indented
    )


def extract_list_items_from_paragraphs(paragraphs: list) -> list:
    """
    Extract list items from paragraph list.
    
    Args:
        paragraphs: List of FormattedParagraph objects
        
    Returns:
        List of FormattedListItem objects
    """
    list_items = []
    
    for paragraph in paragraphs:
        # Get paragraph text
        paragraph_text = ''.join(run.text for run in paragraph.runs)
        
        # Check for list markers
        list_type, level = detect_bullet_or_number(paragraph_text)
        
        if list_type != "none":
            # Remove the list marker from the text
            cleaned_runs = []
            first_run_processed = False
            
            for run in paragraph.runs:
                if not first_run_processed:
                    # Remove marker from first run
                    cleaned_text = re.sub(r'^\s*([•·▪▫■□▶▷◆◇○●★☆-*+]|\d+[.):]|[a-zA-Z][.):]|[ivxlcdmIVXLCDM]+[.):])\s*', '', run.text)
                    if cleaned_text:
                        cleaned_run = FormattedRun(
                            text=cleaned_text,
                            bold=run.bold,
                            italic=run.italic,
                            underline=run.underline,
                            font_name=run.font_name,
                            font_size=run.font_size,
                            color=run.color,
                            highlight=run.highlight
                        )
                        cleaned_runs.append(cleaned_run)
                    first_run_processed = True
                else:
                    cleaned_runs.append(run)
            
            cleaned_paragraph = FormattedParagraph(
                runs=cleaned_runs,
                style=paragraph.style,
                alignment=paragraph.alignment,
                space_before=paragraph.space_before,
                space_after=paragraph.space_after,
                line_spacing=paragraph.line_spacing,
                left_indent=paragraph.left_indent,
                right_indent=paragraph.right_indent,
                first_line_indent=paragraph.first_line_indent
            )
            
            list_item = FormattedListItem(
                paragraph=cleaned_paragraph,
                level=level,
                list_type=list_type,
                number_format="decimal" if list_type == "number" else None
            )
            list_items.append(list_item)
    
    return list_items


def extract_tables_from_page(page) -> list:
    """
    Extract tables from a PDF page.
    
    Args:
        page: pdfplumber page object
        
    Returns:
        List of FormattedTable objects
    """
    tables = []
    
    try:
        # Extract tables using pdfplumber
        page_tables = page.extract_tables()
        
        for table_data in page_tables:
            if not table_data:
                continue
            
            formatted_rows = []
            for row in table_data:
                formatted_cells = []
                for cell in row:
                    if cell is None:
                        cell_text = ""
                    else:
                        cell_text = str(cell).strip()
                    
                    # Create a simple paragraph for each cell
                    if cell_text:
                        cell_run = FormattedRun(text=cell_text)
                        cell_paragraph = FormattedParagraph(runs=[cell_run])
                    else:
                        cell_paragraph = FormattedParagraph(runs=[])
                    
                    formatted_cells.append([cell_paragraph])
                formatted_rows.append(formatted_cells)
            
            if formatted_rows:
                formatted_table = FormattedTable(rows=formatted_rows)
                tables.append(formatted_table)
    
    except Exception:
        # Table extraction failed, return empty list
        pass
    
    return tables


def parse_txt_file(file_or_bytes: Union[str, bytes, io.BytesIO]) -> FormattedDocument:
    """
    Parse a .txt file and build a FormattedDocument.
    
    Args:
        file_or_bytes: File path, bytes, or BytesIO object containing the .txt file
        
    Returns:
        FormattedDocument object containing all document content with basic formatting
        
    Raises:
        Exception: If text parsing fails
    """
    try:
        # Read the text content
        if isinstance(file_or_bytes, str):
            with open(file_or_bytes, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif isinstance(file_or_bytes, bytes):
            text = file_or_bytes.decode('utf-8', errors='ignore')
        elif hasattr(file_or_bytes, 'read'):
            # BytesIO or file-like object
            content = file_or_bytes.read()
            if isinstance(content, bytes):
                text = content.decode('utf-8', errors='ignore')
            else:
                text = content
        else:
            raise Exception("Invalid input type for text file")
        
        # Split into paragraphs and create FormattedDocument
        paragraphs = []
        for line in text.split('\n'):
            if line.strip():  # Only add non-empty lines
                run = FormattedRun(text=line)
                paragraph = FormattedParagraph(runs=[run])
                paragraphs.append(paragraph)
        
        return FormattedDocument(paragraphs=paragraphs, tables=[], list_items=[])
        
    except Exception as e:
        raise Exception(f"Failed to parse text file: {str(e)}")


def text_to_formatted_document(text: str) -> FormattedDocument:
    """
    Convert plain text to a FormattedDocument with basic formatting.
    
    Args:
        text: Plain text string
        
    Returns:
        FormattedDocument object
    """
    paragraphs = []
    for line in text.split('\n'):
        if line.strip():
            run = FormattedRun(text=line)
            paragraph = FormattedParagraph(runs=[run])
            paragraphs.append(paragraph)
    
    return FormattedDocument(paragraphs=paragraphs, tables=[], list_items=[])


def detect_file_type(file_or_bytes) -> str:
    """
    Detect file type from magic bytes and extension.
    
    Args:
        file_or_bytes: File path, bytes, or BytesIO object
        
    Returns:
        File type string ('docx', 'pdf', 'txt', 'unknown')
    """
    magic_bytes = None
    filename = None
    
    if isinstance(file_or_bytes, str):
        # File path
        filename = file_or_bytes
        try:
            with open(file_or_bytes, 'rb') as f:
                magic_bytes = f.read(8)
        except Exception:
            pass
    elif isinstance(file_or_bytes, bytes):
        magic_bytes = file_or_bytes[:8]
    elif hasattr(file_or_bytes, 'read'):
        # BytesIO or file-like object
        current_pos = file_or_bytes.tell()
        file_or_bytes.seek(0)
        magic_bytes = file_or_bytes.read(8)
        file_or_bytes.seek(current_pos)  # Restore position
    
    # Check magic bytes
    if magic_bytes:
        # PDF magic bytes
        if magic_bytes.startswith(b'%PDF'):
            return 'pdf'
        # DOCX magic bytes (ZIP signature)
        elif magic_bytes.startswith(b'PK\x03\x04') or magic_bytes.startswith(b'PK\x05\x06') or magic_bytes.startswith(b'PK\x07\x08'):
            # Could be DOCX, need to check further
            if filename and filename.lower().endswith('.docx'):
                return 'docx'
            elif not filename:  # No filename, assume docx for ZIP files
                return 'docx'
    
    # Fallback to extension
    if filename:
        if filename.lower().endswith('.pdf'):
            return 'pdf'
        elif filename.lower().endswith('.docx'):
            return 'docx'
        elif filename.lower().endswith('.txt'):
            return 'txt'
    
    # If no magic bytes found and no filename, check if it could be txt
    if magic_bytes is None:
        return 'txt'  # Default fallback for plain text
    
    return 'unknown'


def parse_pdf_file(file_or_bytes) -> FormattedDocument:
    """
    Parse a PDF file and build a FormattedDocument.
    
    Args:
        file_or_bytes: File path, bytes, or BytesIO object containing the PDF file
        
    Returns:
        FormattedDocument object containing all document content with formatting
        
    Raises:
        Exception: If PDF parsing fails
    """
    try:
        # Open the PDF
        if isinstance(file_or_bytes, str):
            pdf = pdfplumber.open(file_or_bytes)
        elif isinstance(file_or_bytes, bytes):
            pdf = pdfplumber.open(io.BytesIO(file_or_bytes))
        else:
            pdf = pdfplumber.open(file_or_bytes)
        
        all_paragraphs = []
        all_tables = []
        all_list_items = []
        
        with pdf:
            # Check if PDF is encrypted
            if hasattr(pdf, 'is_encrypted') and pdf.is_encrypted:
                raise Exception("Cannot process encrypted PDF files. Please decrypt the PDF first.")
            
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Extract characters
                    chars = page.chars
                    
                    # Handle scanned PDFs (no extractable text)
                    if not chars:
                        continue  # Skip empty pages, will result in empty document if all pages are scanned
                    
                    # Group characters into paragraphs
                    paragraph_groups = group_text_into_paragraphs(chars)
                    
                    # Convert to FormattedParagraph objects
                    page_paragraphs = []
                    for char_group in paragraph_groups:
                        formatted_paragraph = extract_paragraph_from_chars(char_group)
                        if any(run.text.strip() for run in formatted_paragraph.runs):  # Only add non-empty paragraphs
                            page_paragraphs.append(formatted_paragraph)
                    
                    # Extract list items from paragraphs
                    page_list_items = extract_list_items_from_paragraphs(page_paragraphs)
                    
                    # Remove list items from regular paragraphs
                    list_paragraph_texts = {(''.join(run.text for run in item.paragraph.runs)).strip() for item in page_list_items}
                    non_list_paragraphs = []
                    for para in page_paragraphs:
                        para_text = (''.join(run.text for run in para.runs)).strip()
                        # Check if this paragraph was converted to a list item
                        is_list_paragraph = False
                        for list_text in list_paragraph_texts:
                            if list_text and para_text and (list_text in para_text or para_text in list_text):
                                is_list_paragraph = True
                                break
                        if not is_list_paragraph:
                            non_list_paragraphs.append(para)
                    
                    all_paragraphs.extend(non_list_paragraphs)
                    all_list_items.extend(page_list_items)
                    
                    # Extract tables
                    page_tables = extract_tables_from_page(page)
                    all_tables.extend(page_tables)
                    
                except Exception as e:
                    # Continue processing other pages if one page fails
                    continue
        
        # Create the document
        document = FormattedDocument(
            paragraphs=all_paragraphs,
            tables=all_tables,
            list_items=all_list_items
        )
        
        # If document is completely empty, it might be a scanned PDF
        if (not all_paragraphs and not all_tables and not all_list_items):
            # Return empty document with a warning in the first paragraph
            warning_run = FormattedRun(
                text="Warning: This appears to be a scanned PDF with no extractable text. Content may be images only."
            )
            warning_paragraph = FormattedParagraph(runs=[warning_run])
            document = FormattedDocument(
                paragraphs=[warning_paragraph],
                tables=[],
                list_items=[]
            )
        
        return document
        
    except Exception as e:
        error_message = str(e)
        
        # Provide specific error messages for common issues
        if "encrypted" in error_message.lower():
            raise Exception("Cannot process encrypted PDF files. Please decrypt the PDF first.")
        elif "corrupted" in error_message.lower() or "invalid" in error_message.lower():
            raise Exception(f"PDF file appears to be corrupted or invalid: {error_message}")
        else:
            raise Exception(f"Failed to parse PDF file: {error_message}")


def parse_document_file(file_or_bytes) -> FormattedDocument:
    """
    Parse a document file (TXT, DOCX, or PDF) and build a FormattedDocument.
    
    Args:
        file_or_bytes: File path, bytes, or BytesIO object containing the document
        
    Returns:
        FormattedDocument object containing all document content with formatting
        
    Raises:
        Exception: If document parsing fails or file type is unsupported
    """
    # Detect file type
    file_type = detect_file_type(file_or_bytes)
    
    if file_type == 'txt':
        return parse_txt_file(file_or_bytes)
    elif file_type == 'docx':
        return parse_docx_file(file_or_bytes)
    elif file_type == 'pdf':
        return parse_pdf_file(file_or_bytes)
    else:
        raise Exception(f"Unsupported file type. Only TXT, DOCX, and PDF files are supported.")